import os
import PyPDF2
from docx import Document
from typing import List, Dict, Any
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.schema import Document as LangchainDocument
import logging

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

class DocumentProcessor:
    """Handles document loading and text chunking for RAG system."""
    
    def __init__(self, chunk_size: int = 1000, chunk_overlap: int = 200):
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=chunk_size,
            chunk_overlap=chunk_overlap,
            length_function=len,
            separators=["\n\n", "\n", " ", ""]
        )
    
    def load_pdf(self, file_path: str) -> str:
        """Load text content from PDF file."""
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                text = ""
                for page in pdf_reader.pages:
                    text += page.extract_text() + "\n"
                return text
        except Exception as e:
            logger.error(f"Error loading PDF {file_path}: {str(e)}")
            return ""
    
    def load_docx(self, file_path: str) -> str:
        """Load text content from DOCX file."""
        try:
            doc = Document(file_path)
            text = ""
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            return text
        except Exception as e:
            logger.error(f"Error loading DOCX {file_path}: {str(e)}")
            return ""
    
    def load_txt(self, file_path: str) -> str:
        """Load text content from TXT file."""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                return file.read()
        except Exception as e:
            logger.error(f"Error loading TXT {file_path}: {str(e)}")
            return ""
    
    def load_document(self, file_path: str) -> str:
        """Load document based on file extension."""
        file_extension = os.path.splitext(file_path)[1].lower()
        
        if file_extension == '.pdf':
            return self.load_pdf(file_path)
        elif file_extension == '.docx':
            return self.load_docx(file_path)
        elif file_extension == '.txt':
            return self.load_txt(file_path)
        else:
            logger.warning(f"Unsupported file format: {file_extension}")
            return ""
    
    def chunk_document(self, text: str, metadata: Dict[str, Any] = None) -> List[LangchainDocument]:
        """Split document text into chunks."""
        if not text.strip():
            return []
        
        # Create metadata if not provided
        if metadata is None:
            metadata = {}
        
        # Split text into chunks
        chunks = self.text_splitter.split_text(text)
        
        # Create LangchainDocument objects
        documents = []
        for i, chunk in enumerate(chunks):
            doc_metadata = metadata.copy()
            doc_metadata['chunk_id'] = i
            doc_metadata['chunk_size'] = len(chunk)
            
            documents.append(LangchainDocument(
                page_content=chunk,
                metadata=doc_metadata
            ))
        
        logger.info(f"Created {len(documents)} chunks from document")
        return documents
    
    def process_document(self, file_path: str) -> List[LangchainDocument]:
        """Complete document processing pipeline."""
        # Load document
        text = self.load_document(file_path)
        if not text:
            return []
        
        # Create metadata
        metadata = {
            'source': file_path,
            'filename': os.path.basename(file_path),
            'file_type': os.path.splitext(file_path)[1].lower()
        }
        
        # Chunk document
        return self.chunk_document(text, metadata)
    
    def process_multiple_documents(self, file_paths: List[str]) -> List[LangchainDocument]:
        """Process multiple documents and return all chunks."""
        all_documents = []
        
        for file_path in file_paths:
            if os.path.exists(file_path):
                documents = self.process_document(file_path)
                all_documents.extend(documents)
            else:
                logger.warning(f"File not found: {file_path}")
        
        logger.info(f"Processed {len(file_paths)} files, created {len(all_documents)} total chunks")
        return all_documents
